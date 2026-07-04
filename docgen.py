"""
docgen.py
---------
Takes the fully-executed plan + reflection notes and produces a polished
.docx file using python-docx: proper Heading styles (so it has a real
document outline, not just bold text), a title page block, and a clearly
separated "Assumptions & Notes" section from the reflection step.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)


def build_docx(plan: dict, assumptions: list[str], filename: str) -> str:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block
    title = doc.add_heading(plan.get("title", "Generated Document"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"{plan.get('doc_type', 'Business Document')}  |  Generated {datetime.now().strftime('%B %d, %Y')}"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Sections from the executed plan
    for step in plan["steps"]:
        doc.add_heading(step["section_heading"], level=1)
        content = step.get("content", "")
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(para_text[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(para_text)

    # Reflection / assumptions section
    doc.add_heading("Assumptions & Notes", level=1)
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Generated automatically by the agent's self-check step, based on gaps "
        "or ambiguity in the original request:"
    )
    intro_run.italic = True
    for bullet in assumptions:
        doc.add_paragraph(bullet, style="List Bullet")

    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
